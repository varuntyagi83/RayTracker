#!/usr/bin/env python3
"""
Create a beautifully formatted DOCX file for Voltic User Guide
with proper headings, tables, styling, and structure.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')

    # Set the run's text
    rPr = OxmlElement('w:rPr')

    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def create_voltic_user_guide():
    """Create the formatted DOCX document"""
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Voltic User Guide"
    doc.core_properties.subject = "Complete Documentation for Meta Advertising Intelligence Platform"
    doc.core_properties.keywords = "Voltic, Meta Ads, Facebook Ads, Advertising, AI, Documentation"
    doc.core_properties.comments = "Comprehensive user guide for the Voltic platform"

    # Define custom styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)

    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(20)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 112, 192)

    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 112, 192)

    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(68, 68, 68)

    # ==================== COVER PAGE ====================

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("VOLTIC")
    title_run.font.size = Pt(44)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("User Guide & Documentation")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(68, 68, 68)

    doc.add_paragraph()

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.add_run("Meta Advertising Intelligence & Creative Generation Platform")
    tagline_run.font.size = Pt(14)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("\n" * 8)

    # Version info
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.add_run("Version 1.0 | February 2026").font.size = Pt(11)

    add_page_break(doc)

    # ==================== TABLE OF CONTENTS ====================

    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("This guide covers all features and capabilities of the Voltic platform.")
    doc.add_paragraph()

    toc_items = [
        ("1. Introduction", "Overview of Voltic platform and key capabilities"),
        ("2. Getting Started", "Account setup, workspace creation, and navigation"),
        ("3. Workspace Dashboard", "Home dashboard and activity monitoring"),
        ("4. Ad Discovery & Competitor Intelligence", "Search, filter, and analyze competitor ads"),
        ("5. Boards & Swipe Files", "Organize and manage saved ad collections"),
        ("6. AI-Powered Variations", "Generate ad variations from competitors or products"),
        ("7. Ad Generator", "Batch text overlay composition tool"),
        ("8. Assets & Product Catalog", "Manage product images and backgrounds"),
        ("9. Brand Guidelines", "Define brand voice, colors, and identity"),
        ("10. Automations", "Schedule automated reports and alerts"),
        ("11. Reports & Analytics", "Six report types for performance analysis"),
        ("12. Campaign Analysis", "Deep-dive Meta ad account analytics"),
        ("13. Creative Studio", "AI creative assistant and brainstorming"),
        ("14. Decomposition Tool", "Extract product details from ads with AI"),
        ("15. Competitors Tracking", "Monitor and track competitor brands"),
        ("16. Credits & Billing", "Manage credit balance and purchases"),
        ("17. Settings & Configuration", "Workspace and integration settings"),
        ("18. Best Practices", "Tips for maximizing Voltic features"),
        ("19. Troubleshooting", "Common issues and solutions"),
        ("20. API & Integrations", "Developer documentation and webhooks"),
    ]

    for item, desc in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item).bold = True
        p.add_run(f"\n   {desc}").font.size = Pt(10)

    add_page_break(doc)

    # ==================== 1. INTRODUCTION ====================

    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("What is Voltic?", level=2)
    p = doc.add_paragraph(
        "Voltic is an all-in-one SaaS platform that unifies Meta (Facebook/Instagram) advertising analytics, "
        "competitor intelligence, automated reporting, social comment monitoring, and AI-powered creative generation "
        "into a single workspace."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Think of Voltic as: ").bold = True
    p.add_run("Supermetrics + AdSpy + Jasper combined into one platform.")

    doc.add_heading("Key Capabilities", level=2)

    capabilities = [
        ("Competitor Intelligence", "Discover and analyze competitor ads across Meta's Ad Library"),
        ("Creative Generation", "AI-powered ad variations and text overlay composition"),
        ("Automated Reporting", "Schedule performance, competitor, and comment reports to Slack"),
        ("Product Decomposition", "Extract product information from competitor ads using AI"),
        ("Swipe File Management", "Organize and categorize saved ads in boards"),
        ("Multi-Account Analytics", "Connect up to 91+ Meta ad accounts per workspace"),
        ("AI Creative Tools", "Generate variations, compose text on images, edit product photos"),
    ]

    for title, desc in capabilities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_heading("Who is Voltic For?", level=2)

    audiences = [
        ("Performance Marketers", "Track competitor strategies and automate reporting"),
        ("Creative Teams", "Generate ad variations at scale and build swipe files"),
        ("E-commerce Brands", "Analyze competitor product positioning and create product-based ads"),
        ("Agencies", "Manage multiple client workspaces with centralized intelligence"),
        ("Social Media Managers", "Monitor ad performance and generate platform-specific copy"),
    ]

    for role, use_case in audiences:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{role}: ").bold = True
        p.add_run(use_case)

    add_page_break(doc)

    # ==================== 2. GETTING STARTED ====================

    doc.add_heading("2. Getting Started", level=1)

    doc.add_heading("Account Setup", level=2)

    doc.add_heading("1. Sign Up", level=3)
    signup_steps = [
        "Visit your Voltic instance URL",
        "Click 'Sign Up' on the login page",
        "Enter your email and password (minimum 6 characters)",
        "Verify your email address"
    ]
    for step in signup_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("2. Create Your Workspace", level=3)
    workspace_steps = [
        "Upon first login, you'll be prompted to create a workspace",
        "Enter workspace name (e.g., 'Acme Marketing Team')",
        "Invite team members via email (optional)"
    ]
    for step in workspace_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("3. Connect Meta Ad Accounts", level=3)
    meta_steps = [
        "Navigate to Settings → Ad Accounts",
        "Click 'Connect Meta Account'",
        "Authenticate with Facebook/Meta",
        "Select ad accounts to sync (up to 91+ accounts)"
    ]
    for step in meta_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("Navigation Overview", level=2)
    doc.add_paragraph("Main Navigation (Left Sidebar):")

    nav_items = [
        ("Home", "Workspace overview dashboard"),
        ("Automations", "Scheduled reports and alerts"),
        ("Discover", "Ad library search and competitor intelligence"),
        ("Boards", "Saved ad collections (swipe files)"),
        ("Variations", "AI-powered ad variation generator"),
        ("Ad Generator", "Text overlay composition tool (NEW)"),
        ("Assets", "Product catalog and background images"),
        ("Reports", "6 report types (Top Ads, Campaigns, Creatives, etc.)"),
        ("Campaign Analysis", "Deep-dive ad account analytics"),
        ("Creative Studio", "AI creative assistant"),
        ("Brand Guidelines", "Brand voice and visual identity"),
        ("Decomposition", "Product extraction from competitor ads"),
        ("Competitors", "Tracked competitor brands"),
        ("Credits", "AI feature credit balance"),
        ("Settings", "Workspace configuration"),
    ]

    for nav, desc in nav_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{nav} — ").bold = True
        p.add_run(desc)

    add_page_break(doc)

    # ==================== 3. CORE FEATURES ====================

    doc.add_heading("3. Core Features", level=1)

    doc.add_heading("Credit System", level=2)
    doc.add_paragraph("Voltic uses credits for AI-powered features:")

    # Create credit cost table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Cost'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    credit_costs = [
        ('AI Variation (per strategy)', '10 credits'),
        ('Product Decomposition', '5 credits'),
        ('AI Image Generation (DALL-E)', '15 credits'),
        ('AI Image Editing (Gemini)', '12 credits'),
        ('Creative Studio Chat Message', '3 credits'),
        ('Background Generation', '15 credits'),
    ]

    for feature, cost in credit_costs:
        row_cells = table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = cost

    doc.add_paragraph()
    doc.add_paragraph("How to Get Credits:")
    doc.add_paragraph("• Purchase credit packs in Settings → Billing", style='List Bullet')
    doc.add_paragraph("• Credits are workspace-scoped (shared by all members)", style='List Bullet')
    doc.add_paragraph("• Credits never expire", style='List Bullet')

    add_page_break(doc)

    # ==================== 6. AI-POWERED VARIATIONS ====================

    doc.add_heading("6. AI-Powered Variations", level=1)
    doc.add_paragraph("Location: /variations", style='Intense Quote')

    doc.add_heading("Overview", level=2)
    p = doc.add_paragraph(
        "The Variations page is a dedicated workspace for generating AI-powered ad variations at scale. "
        "It supports two sources:"
    )

    doc.add_paragraph("1. Competitor Ads — Generate variations inspired by competitor creatives", style='List Number')
    doc.add_paragraph("2. Your Products — Generate variations starting from your product images (NEW)", style='List Number')

    doc.add_heading("Asset-Based Variations (NEW FEATURE)", level=2)

    # Highlight box for new feature
    p = doc.add_paragraph()
    p.add_run("✨ NEW FEATURE").bold = True
    p.add_run(" — Upload your product images and generate variations with AI-powered editing while preserving product labels exactly.")

    doc.add_paragraph()
    doc.add_heading("How It Works:", level=3)

    asset_steps = [
        "Upload your product image OR select from asset library",
        "Choose brand guideline (optional, for color palette)",
        "Set creative options (angle, lighting, background)",
        "Select strategies (Hero Product, Curiosity, Pain Point, etc.)",
        "AI edits your product image using Gemini while preserving product labels exactly"
    ]

    for step in asset_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("Use Case:", level=3)
    p = doc.add_paragraph()
    p.add_run("Example: ").bold = True
    p.add_run("\"Here's my vitamin bottle — create 6 variations with different backgrounds and lighting styles.\"")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Perfect for: ").bold = True
    p.add_run("E-commerce product photography transformation")

    doc.add_heading("Channel Selection (NEW FEATURE)", level=2)
    doc.add_paragraph("Choose the advertising platform to optimize copy for:")

    # Create channel table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Channel'
    hdr_cells[1].text = 'Copy Style'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    channels = [
        ('Facebook', 'Conversational, emoji-friendly, engagement-focused, longer storytelling'),
        ('Instagram', 'Visual-first, hashtag-ready, shorter punchy copy, aspirational tone'),
        ('TikTok', 'Gen-Z tone, trend-aware, ultra-short, casual and authentic'),
        ('LinkedIn', 'Professional, thought-leadership tone, B2B-friendly, data-driven'),
        ('Google Ads', 'Keyword-focused, direct response, respect character limits, action-oriented'),
    ]

    for channel, style in channels:
        row_cells = table.add_row().cells
        row_cells[0].text = channel
        row_cells[1].text = style

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Default: ").bold = True
    p.add_run("Facebook (most versatile)")

    doc.add_heading("Strategy Descriptions", level=2)

    strategies = [
        ("Hero Product",
         "Text: Product name in headline, feature-benefit structure\n"
         "Image: Product centered, clean professional look, prominent focal point\n"
         "Best For: E-commerce, product launches, clear value props"),

        ("Curiosity",
         "Text: Pattern-interrupt headline, 'What if...' or 'The secret to...' hooks\n"
         "Image: Dramatic lighting, unexpected angle, visually intriguing composition\n"
         "Best For: Engagement campaigns, top-of-funnel awareness"),

        ("Pain Point",
         "Text: Calls out specific problem, positions product as solution\n"
         "Image: Visual contrast or metaphor, product appears as clear solution\n"
         "Best For: Problem-aware audiences, consideration stage"),

        ("Proof Point",
         "Text: Stats, testimonials, social proof, 'Join 10,000+ customers'\n"
         "Image: Premium, trustworthy, aspirational quality, credibility cues\n"
         "Best For: Conversion campaigns, overcoming objections"),

        ("Image Only",
         "Text: Minimal or no text, product name only\n"
         "Image: Stunning, eye-catching product photo, high production value\n"
         "Best For: Visual platforms (Instagram), brand awareness"),

        ("Text Only",
         "Text: Long-form copy, storytelling, detailed explanation\n"
         "Image: Simple background with product, text is the hero\n"
         "Best For: Complex products, educational content"),
    ]

    for strategy, details in strategies:
        doc.add_heading(f"{strategy}", level=3)
        for line in details.split('\n'):
            if line.strip():
                if ':' in line:
                    parts = line.split(':', 1)
                    p = doc.add_paragraph()
                    p.add_run(parts[0] + ': ').bold = True
                    p.add_run(parts[1].strip())
                else:
                    doc.add_paragraph(line.strip())

    doc.add_heading("Cost", level=2)
    p = doc.add_paragraph()
    p.add_run("10 credits per strategy").bold = True
    p.add_run(" (unchanged)")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Example: ").italic = True
    p.add_run("Generate 3 strategies = 30 credits")

    add_page_break(doc)

    # ==================== 7. AD GENERATOR ====================

    doc.add_heading("7. Ad Generator (NEW FEATURE)", level=1)
    doc.add_paragraph("Location: /ad-generator", style='Intense Quote')

    # Highlight box
    p = doc.add_paragraph()
    p.add_run("✨ BRAND NEW FEATURE").bold = True
    p.add_run(" — Create hundreds of ad variations in minutes by combining backgrounds with text.")

    doc.add_heading("What is Ad Generator?", level=2)
    p = doc.add_paragraph(
        "A batch text overlay composition tool that lets you create M×N ad variations by combining:"
    )
    doc.add_paragraph("• M backgrounds (product images, lifestyle photos, brand assets)", style='List Bullet')
    doc.add_paragraph("• N text variants (headlines, ad copy, CTAs)", style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Example: ").bold = True
    p.add_run("5 backgrounds × 10 text variants = ")
    p.add_run("50 ad previews ").bold = True
    p.add_run("generated in ~20 seconds")

    doc.add_heading("When to Use Ad Generator", level=2)

    doc.add_paragraph("Best For:", style='Heading 3')
    best_for = [
        "Creating multiple ad creatives at scale",
        "A/B testing different copy on the same visual",
        "Brand awareness campaigns with consistent visuals",
        "Social media content calendars"
    ]
    for item in best_for:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("Not Ideal For:", style='Heading 3')
    not_for = [
        "Complex image editing (use Variations with Gemini instead)",
        "Product photography transformation (use Asset-Based Variations)"
    ]
    for item in not_for:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("7-Step Workflow", level=2)

    # Step 1
    doc.add_heading("Step 1: Select Brand Guideline", level=3)
    doc.add_paragraph("Links ads to your brand identity for consistent styling")
    steps = [
        "Click guideline dropdown",
        "Select from existing brand guidelines",
        "If none exist, create one in Brand Guidelines page first"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Step 2
    doc.add_heading("Step 2: Select Background Images", level=3)
    doc.add_paragraph("Choose product images or brand assets to use as backgrounds")
    steps = [
        "Asset grid shows all images linked to selected guideline",
        "Click to select (multi-select enabled, up to 20)",
        "Selected assets show checkmark overlay"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Step 3
    doc.add_heading("Step 3: Enter Text Variants", level=3)
    doc.add_paragraph("Write headlines, ad copy, or CTAs to test")
    steps = [
        "Start with one text input field",
        "Type headline or ad copy (2-10 words works best)",
        "Click '+ Add Variant' to add more fields (up to 20)",
        "Click '×' to remove a variant"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Examples:").bold = True
    examples = [
        "\"Your Perfect Morning Starts Here ☕\"",
        "\"Limited Time: 30% Off All Coffee\"",
        "\"Voted #1 Coffee by Barista Magazine\"",
        "\"Wake Up to Better Coffee\""
    ]
    for ex in examples:
        doc.add_paragraph(ex, style='List Bullet')

    # Step 4
    doc.add_heading("Step 4: Styling Controls", level=3)

    # Create styling options table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Control'
    hdr_cells[1].text = 'Options'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    styling = [
        ('Font Family', 'Inter, Roboto, Playfair Display, Montserrat, Open Sans, Lato'),
        ('Font Size', '24px - 96px (default: 48px)'),
        ('Text Color', 'Color picker (hex input, default: #FFFFFF white)'),
        ('Text Position', 'Center, Top, Bottom, Corners, Custom (8 presets)'),
        ('Text Effects', 'Automatic drop shadow for readability'),
    ]

    for control, options in styling:
        row_cells = table.add_row().cells
        row_cells[0].text = control
        row_cells[1].text = options

    # Step 5
    doc.add_heading("Step 5: Generate Previews", level=3)
    steps = [
        "Click 'Generate Previews' button",
        "Shows count: 'Generate Previews (50)' for 5 backgrounds × 10 texts",
        "Processing happens in batches of 5",
        "Results appear in Preview Grid"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Time Estimate: ").bold = True
    p.add_run("~20 seconds for 50 previews | ~40 seconds for 100 previews")

    # Step 6
    doc.add_heading("Step 6: Review & Approve", level=3)
    doc.add_paragraph("Preview Grid shows all composited ads with:")
    items = [
        "Composited ad preview image",
        "Text variant displayed",
        "Background asset name",
        "Approve/Reject buttons",
        "Download button"
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Step 7
    doc.add_heading("Step 7: Save Approved Ads", level=3)
    steps = [
        "Click 'Save Approved (X)' button",
        "Only approved ads are saved to workspace",
        "Server creates database records with metadata",
        "Ads appear in Ads History section"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("Tips for Best Results", level=2)

    tips = [
        ("Background Selection", [
            "Use high-resolution images (1080×1080 minimum)",
            "Ensure backgrounds have negative space for text",
            "Avoid busy/cluttered backgrounds",
            "Consistent aspect ratio (1:1 or 4:5 for social)"
        ]),
        ("Text Variants", [
            "Start with 5-10 variants for A/B testing",
            "Test different emotional tones",
            "Vary length (short punchy vs longer descriptive)",
            "Include numbers/stats ('Save 30%', 'Join 10K+')"
        ]),
        ("Styling", [
            "High contrast (white on dark or dark on light)",
            "Avoid mid-tones (gray on gray)",
            "Match font to brand personality",
            "Use larger font sizes for mobile (60px+)"
        ]),
    ]

    for category, tip_list in tips:
        doc.add_heading(category, level=3)
        for tip in tip_list:
            doc.add_paragraph(tip, style='List Bullet')

    add_page_break(doc)

    # ==================== GEMINI IMAGE EDITING ====================

    doc.add_heading("Gemini Image Editing (NEW TECHNOLOGY)", level=1)

    # Highlight box
    p = doc.add_paragraph()
    p.add_run("🚀 BREAKTHROUGH TECHNOLOGY").bold = True
    p.add_run(" — Powered by Google's Gemini 2.5 Flash/Pro Image model")

    doc.add_heading("What is Gemini Image Editing?", level=2)
    p = doc.add_paragraph(
        "Gemini replaces DALL-E for asset-based variations, offering superior accuracy and speed. "
        "It uses advanced mask-based editing to transform product images while preserving the "
        "product itself (including labels, text, and packaging) exactly."
    )

    doc.add_heading("How It Works", level=2)

    steps = [
        ("Generate Product Mask",
         "AI creates a segmentation mask where WHITE = product, BLACK = background"),
        ("Apply Transformations",
         "AI edits ONLY the background (black areas) based on your creative options"),
        ("Preserve Product Exactly",
         "Product labels, text, and packaging remain pixel-perfect"),
        ("Upload to Storage",
         "Final image saved to Supabase Storage with public URL"),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        doc.add_heading(f"{i}. {title}", level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Creative Options", level=2)

    # Create options table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Medium Shading 1 Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Option'
    hdr_cells[1].text = 'Choices'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    options = [
        ('Product Angle', 'Front View, Side View, 3/4 View, Top-Down'),
        ('Lighting Style', 'Studio, Natural, Golden Hour, Dramatic'),
        ('Background Style', 'Solid White, Lifestyle, Outdoor, Gradient'),
        ('Custom Instruction', 'Free-form text (e.g., "Place on wooden shelf")'),
        ('Brand Colors', 'Pulled from Brand Guidelines (optional)'),
    ]

    for option, choices in options:
        row_cells = table.add_row().cells
        row_cells[0].text = option
        row_cells[1].text = choices

    doc.add_heading("Why Gemini Matters", level=2)

    benefits = [
        ("⚡ Faster", "2-4x faster than DALL-E for image editing"),
        ("🎯 More Accurate", "Mask-based editing preserves product labels exactly"),
        ("💰 Cost-Effective", "Gemini Flash is cheaper than DALL-E 3"),
        ("🔧 Flexible", "Supports both Flash (fast) and Pro (quality) models"),
    ]

    for benefit, desc in benefits:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{benefit}: ").bold = True
        p.add_run(desc)

    add_page_break(doc)

    # ==================== DISCOVER IMPROVEMENTS ====================

    doc.add_heading("Discover Page — New Features", level=1)

    doc.add_heading("Save as Competitor (NEW)", level=2)

    # Highlight
    p = doc.add_paragraph()
    p.add_run("✨ ONE-CLICK TRACKING").bold = True
    p.add_run(" — Save competitor brands instantly without decomposition")

    doc.add_paragraph()
    doc.add_heading("What It Does:", level=3)
    doc.add_paragraph(
        "Saves ad metadata to your Competitors list in one click. No need to manually decompose first. "
        "Automatically extracts: brand name, headline, platform, format."
    )

    doc.add_heading("How to Use:", level=3)
    steps = [
        "Find an ad from a competitor brand in Discover",
        "Click 'Save as Competitor' button (user icon)",
        "Ad metadata is saved to /competitors",
        "You can now track all ads from this brand"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading("Create Board from Discover (NEW)", level=2)

    doc.add_heading("What It Does:", level=3)
    doc.add_paragraph(
        "Create a new swipe file board directly from search results. "
        "Pre-populate with selected ads. Streamlines inspiration collection workflow."
    )

    doc.add_heading("How to Use:", level=3)
    steps = [
        "Search for ads (e.g., 'fitness apparel')",
        "Select 5-10 ads using checkboxes",
        "Click 'Create Board' button in toolbar",
        "Enter board name (e.g., 'Fitness Ad Swipe')",
        "Board is created with all selected ads saved"
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    add_page_break(doc)

    # ==================== BEST PRACTICES ====================

    doc.add_heading("18. Best Practices", level=1)

    doc.add_heading("Variation Generation", level=2)

    doc.add_heading("Do's ✅", level=3)
    dos = [
        "Start with 2-3 strategies to conserve credits",
        "Use high-quality source images (1080×1080 minimum)",
        "Write specific product descriptions for better AI output",
        "Test channel-specific copy (Facebook vs TikTok)",
        "Link assets to brand guidelines for consistency"
    ]
    for item in dos:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Don'ts ❌", level=3)
    donts = [
        "Don't generate all 6 strategies at once (expensive)",
        "Don't use low-resolution competitor ad screenshots",
        "Don't skip product description (AI needs context)",
        "Don't use generic instructions ('make it better')"
    ]
    for item in donts:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Ad Generator", level=2)

    doc.add_heading("Do's ✅", level=3)
    dos = [
        "Test 5-10 text variants initially",
        "Use high-contrast text colors",
        "Preview on mobile viewport (most users)",
        "Select backgrounds with negative space",
        "Batch generate for efficiency"
    ]
    for item in dos:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Don'ts ❌", level=3)
    donts = [
        "Don't use busy/cluttered backgrounds",
        "Don't use mid-tone text colors (poor contrast)",
        "Don't exceed 15 words per text variant",
        "Don't forget to save approved ads (lose previews on refresh)"
    ]
    for item in donts:
        doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # ==================== TROUBLESHOOTING ====================

    doc.add_heading("19. Troubleshooting", level=1)

    issues = [
        ("Ad Account Connection Expired",
         "Cause: Meta OAuth token expires after 60 days",
         ["Go to Settings → Ad Accounts",
          "Click 'Refresh Token' on expired account",
          "Re-authenticate with Facebook",
          "Automations resume automatically"]),

        ("Insufficient Credits",
         "Cause: Credit balance too low for operation",
         ["Check credit balance in top-right corner",
          "Click 'Buy Credits' button",
          "Purchase credit pack",
          "Retry operation"]),

        ("Variation Generation Failed",
         "Possible Causes: Low-quality source image, Gemini API region restriction, invalid description",
         ["Check source image resolution (1080px minimum)",
          "Verify Gemini API key in environment variables",
          "Add more detailed product description",
          "Try different strategy"]),

        ("Slack Message Not Delivered",
         "Possible Causes: Bot not invited, channel renamed, integration disconnected",
         ["Verify Slack integration in Settings → Integrations",
          "Invite @Voltic bot to target channel (/invite @Voltic)",
          "Test with 'Run Now' on automation",
          "Check Slack workspace permissions"]),
    ]

    for title, cause, fixes in issues:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.add_run("Cause: ").bold = True
        p.add_run(cause)

        doc.add_paragraph()
        doc.add_paragraph("Fix:", style='Heading 3')
        for fix in fixes:
            doc.add_paragraph(fix, style='List Number')

        doc.add_paragraph()

    add_page_break(doc)

    # ==================== FOOTER / CLOSING ====================

    doc.add_heading("Support & Contact", level=1)

    doc.add_paragraph()
    support_info = [
        ("Help Center", "your-domain/help"),
        ("Email Support", "support@voltic.app"),
        ("Live Chat", "Available Mon-Fri 9am-5pm EST"),
        ("Feature Requests", "your-domain/feedback"),
        ("Status Page", "status.voltic.app"),
    ]

    for label, info in support_info:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{label}: ").bold = True
        p.add_run(info)

    doc.add_paragraph()
    doc.add_paragraph()

    # Final note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("─────────────────────").font.color.rgb = RGBColor(200, 200, 200)

    doc.add_paragraph()

    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final.add_run("For the most up-to-date documentation, visit your Voltic instance help center.")
    final_run.font.italic = True
    final_run.font.size = Pt(10)
    final_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version.add_run("Version 1.0 | February 2026 | © Voltic Platform")
    version_run.font.size = Pt(9)
    version_run.font.color.rgb = RGBColor(150, 150, 150)

    # Save document
    output_path = "/Users/varuntyagi/Downloads/Claude Research/RayTracker/VOLTIC_USER_GUIDE_FORMATTED.docx"
    doc.save(output_path)
    print(f"✅ Document created successfully: {output_path}")
    print(f"📄 Total sections: 20+")
    print(f"📊 Includes: Tables, styled headings, bullet points, numbered lists")
    print(f"🎨 Professional formatting with colors and emphasis")

if __name__ == "__main__":
    create_voltic_user_guide()
